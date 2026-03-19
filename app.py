import streamlit as st
import requests
import pandas as pd
import html
from io import BytesIO
import datetime
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from bs4 import BeautifulSoup
import calendar
import time
import re
from dateutil import parser

# ==========================================
# CONFIGURACIÓN INICIAL Y ESTILOS
# ==========================================
st.set_page_config(page_title="Boletín Mensual", layout="wide")

st.markdown("""
    <style>
    div.stButton > button, div.stDownloadButton > button {
        background-color: #00205B !important;
        color: white !important;
        border: none !important;
    }
    div.stButton > button:hover, div.stDownloadButton > button:hover {
        background-color: #00153D !important;
        color: white !important;
    }
    span[data-baseweb="tag"] {
        background-color: #00205B !important;
        color: white !important;
    }
    .github-footer {
        position: fixed;
        right: 20px;
        bottom: 20px;
        background-color: rgba(255, 255, 255, 0.9);
        padding: 8px 12px;
        border-radius: 50px;
        border: 1px solid #d0d7de;
        z-index: 1000;
        display: flex;
        align-items: center;
        font-family: 'Calibri', sans-serif;
        text-decoration: none;
        color: #24292f;
        box-shadow: 0px 4px 12px rgba(0,0,0,0.1);
        transition: transform 0.2s, box-shadow 0.2s;
    }
    .github-footer:hover {
        transform: translateY(-2px);
        box-shadow: 0px 6px 16px rgba(0,0,0,0.15);
        color: #00205B;
        border-color: #00205B;
    }
    .github-icon {
        margin-right: 8px;
        width: 22px;
        height: 22px;
    }
    </style>
    <a class="github-footer" href="https://github.com/sdiazprado" target="_blank">
        <img class="github-icon" src="https://github.githubassets.com/images/modules/logos_page/GitHub-Mark.png" alt="GitHub Logo">
        <span><strong>@sdiazprado</strong></span>
    </a>
""", unsafe_allow_html=True)

# ==========================================
# UTILIDADES DE FORMATO
# ==========================================
def clean_author_name(name):
    if not name:
        return ""
    cleaned = name.strip().title()
    cleaned = re.sub(r'\b([A-Z])\.\s*([A-Z])', lambda m: f"{m.group(1)}. {m.group(2)}", cleaned)
    return cleaned

# ==========================================
# FUNCIONES DE EXTRACCIÓN (BACKEND)
# ==========================================

# --- SECCIÓN: REPORTES ---
# BID (Annual Reports en inglés)
@st.cache_data(show_spinner=False)
def load_reportes_bid_en(start_date_str, end_date_str):
    """
    Extrae Annual Reports del BID en inglés
    URL: https://publications.iadb.org/en?f%5B0%5D=type%3AAnnual%20Reports
    """
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import pandas as pd
    import time
    import re
    from dateutil import parser

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        print(f"📅 Rango de fechas: {start_date.date()} a {end_date.date()}")
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()
        print(f"⚠️ Error en fechas, usando rango por defecto")

    rows = []
    
    # Configuración de paginación
    page = 0
    max_pages = 5  # Límite de páginas a extraer
    hay_resultados = True
    
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
    chrome_options.add_argument("--disable-blink-features=AutomationControlled")
    chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
    chrome_options.add_experimental_option('useAutomationExtension', False)

    try:
        print("🔍 Iniciando Selenium para BID Annual Reports (EN)...")
        driver = webdriver.Chrome(options=chrome_options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        while page < max_pages and hay_resultados:
            # URL para Annual Reports en inglés
            url = f"https://publications.iadb.org/en?f%5B0%5D=type%3AAnnual%20Reports&page={page}"
            
            print(f"\n📄 Accediendo a página {page+1}: {url}")
            driver.get(url)

            try:
                WebDriverWait(driver, 20).until_not(
                    EC.title_contains("Just a moment")
                )
                print(f"✅ Página {page+1} cargada correctamente.")
            except:
                print(f"⚠️ La página {page+1} sigue mostrando 'Just a moment...', esperando...")
                time.sleep(10)

            time.sleep(5)
            html = driver.page_source
            soup = BeautifulSoup(html, 'html.parser')

            # Guardar HTML para depuración (solo primera página)
            if page == 0:
                with open("bid_reportes_debug.html", "w", encoding="utf-8") as f:
                    f.write(html)
                print("💾 HTML guardado en bid_reportes_debug.html")

            # Estrategias de búsqueda
            items = soup.find_all('div', class_='views-row')
            print(f"📚 Página {page+1} - Elementos encontrados: {len(items)}")

            if len(items) == 0:
                print(f"📭 No hay más elementos en página {page+1}")
                hay_resultados = False
                break

            # Mapeo de meses en inglés
            meses_en = {
                'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
                'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
            }

            docs_en_pagina = 0
            for idx, item in enumerate(items):
                print(f"\n--- Procesando elemento {idx+1} ---")
                
                # ESTRATEGIA 1: Buscar específicamente el div con clase 'views-field-field-title'
                title_elem = None
                title_container = item.find('div', class_='views-field-field-title')
                if title_container:
                    span_field = title_container.find('span', class_='field-content')
                    if span_field:
                        a_tag = span_field.find('a')
                        if a_tag:
                            title_elem = a_tag
                            print(f"  ✅ Título encontrado con estrategia 1")

                # ESTRATEGIA 2: Buscar span.field-content > a (estructura genérica)
                if not title_elem:
                    span_field = item.find('span', class_='field-content')
                    if span_field:
                        a_tag = span_field.find('a')
                        if a_tag:
                            title_elem = a_tag
                            print(f"  ✅ Título encontrado con estrategia 2")

                # ESTRATEGIA 3: Buscar cualquier enlace con texto largo
                if not title_elem:
                    for a_tag in item.find_all('a', href=True):
                        texto = a_tag.get_text(strip=True)
                        if len(texto) > 30:
                            title_elem = a_tag
                            print(f"  ✅ Título encontrado con estrategia 3")
                            break

                if not title_elem:
                    print(f"  ⚠️ No se encontró título en elemento")
                    continue

                titulo = title_elem.get_text(strip=True)
                link = title_elem['href']
                if not link.startswith('http'):
                    link = "https://publications.iadb.org" + link

                print(f"  📌 Título extraído: '{titulo[:100]}...'")

                # Extraer fecha - VERSIÓN MEJORADA
                parsed_date = None
                
                # Buscar específicamente el contenedor de fecha
                date_container = item.find('div', class_='views-field-field-date-issued-text')
                if date_container:
                    date_span = date_container.find('span', class_='field-content')
                    if date_span:
                        date_text = date_span.get_text(strip=True)
                        print(f"  📅 Texto de fecha (específico): {date_text}")
                        
                        # Intentar parsear con regex (ej: "Mar 2026")
                        match = re.search(r'([A-Za-z]{3,9})\s+(\d{4})', date_text)
                        if match:
                            mes_str, año_str = match.groups()
                            mes_num = meses_en.get(mes_str.lower()[:3])
                            if mes_num:
                                parsed_date = datetime.datetime(int(año_str), mes_num, 1)
                                print(f"  ✅ Fecha parseada: {parsed_date}")
                
                # Fallback: buscar cualquier span con texto de fecha
                if not parsed_date:
                    for span in item.find_all('span'):
                        text = span.get_text(strip=True)
                        match = re.search(r'([A-Za-z]{3,9})\s+(\d{4})', text)
                        if match:
                            mes_str, año_str = match.groups()
                            mes_num = meses_en.get(mes_str.lower()[:3])
                            if mes_num:
                                parsed_date = datetime.datetime(int(año_str), mes_num, 1)
                                print(f"  ✅ Fecha parseada (fallback): {parsed_date}")
                                break

                if not parsed_date:
                    print(f"  ⚠️ No se pudo extraer fecha")
                    continue

                print(f"  📅 Fecha final: {parsed_date.date()}")

                # Filtrar por fecha
                if parsed_date < start_date or parsed_date > end_date:
                    print(f"  ⏭️ Fecha fuera de rango: {parsed_date.date()} (rango: {start_date.date()} a {end_date.date()})")
                    continue

                # Evitar duplicados
                if not any(r['Link'] == link for r in rows):
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "BID (Reportes)"
                    })
                    docs_en_pagina += 1
                    print(f"  ✅ Documento AGREGADO: {titulo[:50]}...")

            print(f"\n📊 Documentos agregados en esta página: {docs_en_pagina}")
            print(f"📊 Total documentos hasta ahora: {len(rows)}")

            page += 1
            print(f"➡️ Avanzando a página {page+1}...\n")

        driver.quit()

    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ Documentos BID (Reportes) encontrados en {page} páginas: {len(df)}")
        print("\n📋 Primeros documentos:")
        for i, row in df.head(3).iterrows():
            print(f"  - {row['Date'].strftime('%Y-%m')}: {row['Title'][:80]}...")
    else:
        print("\n⚠️ No se encontraron documentos del BID (Reportes)")

    return df

@st.cache_data(show_spinner=False)
def load_reportes_bpi(start_date_str, end_date_str):
    urls_api = [
        "https://www.bis.org/api/document_lists/bcbspubls.json",
        "https://www.bis.org/api/document_lists/cpmi_publs.json"
    ]
    urls_html = ["https://www.bis.org/ifc/publications.htm"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            lista_documentos = data.get("list", {})
            for path, doc_info in lista_documentos.items():
                titulo = html.unescape(doc_info.get("short_title", ""))
                if not titulo: continue
                link = "https://www.bis.org" + doc_info.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"):
                    link += ".htm"
                date_str = doc_info.get("publication_start_date", "")
                parsed_date = None
                if date_str:
                    try: parsed_date = parser.parse(date_str)
                    except: pass
                if not parsed_date: continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue

    for url in urls_html:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            content_div = soup.find('div', id='cmsContent')
            if not content_div: continue
            for p in content_div.find_all('p'):
                a_tag = p.find('a')
                if not a_tag: continue
                titulo = a_tag.get_text(strip=True)
                href = a_tag.get('href', '')
                if not href or 'index.htm' in href: continue 
                link = "https://www.bis.org" + href if href.startswith('/') else href
                full_text = p.get_text(strip=True)
                date_str = full_text.replace(titulo, '').strip(', ')
                parsed_date = None
                if date_str:
                    try: parsed_date = parser.parse(date_str)
                    except: pass
                if not parsed_date:
                    match = re.search(r'\b(20\d{2})\b', titulo)
                    if match: parsed_date = datetime.datetime(int(match.group(1)), 1, 1)
                if not parsed_date: continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "BPI"})
        except Exception as e:
            continue
            
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_reportes_bm(start_date_str, end_date_str):
    """Extractor para Reportes del BM (Solo incluye los que mencionan 'Report')"""
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}
    
    # ID exacto de la comunidad compartida con Investigación
    scope_id = '06251f8a-62c2-59fb-add5-ec0993fc20d9'
    
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows, page = [], 0
    while True:
        try:
            params = {
                'scope': scope_id, 
                'sort': 'dc.date.issued,DESC', 
                'page': page, 
                'size': 20
            }
            res = requests.get(base_url, headers=headers, params=params, timeout=15)
            data = res.json()
            
            objects = data.get('_embedded', {}).get('searchResult', {}).get('_embedded', {}).get('objects', [])
            if not objects: break
            
            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})
                
                # Extraer Título y Fecha (Sin Autor, como acordamos)
                title = meta.get('dc.title', [{'value': ''}])[0].get('value', '')
                date_s = meta.get('dc.date.issued', [{'value': ''}])[0].get('value', '')
                
                parsed_date = None
                if date_s:
                    try: parsed_date = parser.parse(date_s)
                    except: pass
                
                if not parsed_date or parsed_date < start_date: continue
                
                # --- NUEVO FILTRO PRO-REPORTES ---
                abstract_list = meta.get('dc.description.abstract', [])
                desc_list = meta.get('dc.description', [])
                
                description = ""
                if abstract_list: description = abstract_list[0].get('value', '').lower()
                elif desc_list: description = desc_list[0].get('value', '').lower()
                
                # Si la palabra "report" NO está en la descripción, lo saltamos
                if not re.search(r'\breport\b', description):
                    continue
                # ----------------------------------
                
                # Link permanente
                link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get('value', '')
                if not link: link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"
                
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": title, "Link": link, "Organismo": "BM"})
                    items_found += 1
            
            if items_found == 0: break
            page += 1
            if page > 3: break # Límite para evitar búsquedas infinitas
            time.sleep(0.2)
        except:
            break
            
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_reportes_cef(start_date_str, end_date_str):
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        url = f"https://www.fsb.org/publications/?dps_paged={page}"
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            items = soup.find_all('div', class_=lambda c: c and 'post-excerpt' in c)
            if not items: break
            items_found = 0
            for item in items:
                title_div = item.find('div', class_='post-title')
                if not title_div or not title_div.find('a'): continue
                a_tag = title_div.find('a')
                titulo_raw = a_tag.get_text(strip=True)
                link = a_tag.get('href', '')
                date_div = item.find('div', class_='post-date')
                parsed_date = None
                if date_div:
                    try: parsed_date = parser.parse(date_div.get_text(strip=True))
                    except: pass
                if not parsed_date: continue
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "CEF"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.5) 
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_reportes_ocde(start_date_str, end_date_str):
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    rows = []
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    year = start_date.year
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    try:
        driver = webdriver.Chrome(options=chrome_options)
        url = f"https://www.oecd.org/en/search/publications.html?orderBy=mostRecent&page=0&facetTags=oecd-content-types%3Apublications%2Freports%2Coecd-languages%3Aen&minPublicationYear={year}&maxPublicationYear={year}"
        driver.get(url)
        time.sleep(12) 
        js_script = """
        let linksData = [];
        function findLinks(root) {
            let els = root.querySelectorAll('*');
            els.forEach(el => {
                if (el.shadowRoot) findLinks(el.shadowRoot);
                if (el.tagName === 'A' && el.href) {
                    let text = el.innerText || el.textContent;
                    let aria = el.getAttribute('aria-label') || el.getAttribute('title') || '';
                    let final_text = text.trim() ? text.trim() : aria.trim();
                    if(final_text.length > 15) { linksData.push({ title: final_text, link: el.href }); }
                }
            });
        }
        findLinks(document); return linksData;
        """
        extracted_links = driver.execute_script(js_script)
        driver.quit()
        for item in extracted_links:
            href = item['link'].lower()
            title = item['title'].replace('\n', ' ')
            firmas_validas = ['/publications/', '/reports/', 'oecd-ilibrary.org', '/books/']
            if any(firma in href for firma in firmas_validas):
                if any(basura in title.lower() for x in ['download', 'read more', 'pdf', 'buy', 'search', 'subscribe']): continue
                if not any(r['Link'] == item['link'] for r in rows):
                    rows.append({"Date": start_date, "Title": title, "Link": item['link'], "Organismo": "OCDE"})
    except: pass
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df


@st.cache_data(show_spinner=False)
def load_reportes_bpi(start_date_str, end_date_str):
    urls_api = ["https://www.bis.org/api/document_lists/bcbspubls.json", "https://www.bis.org/api/document_lists/cpmi_publs.json"]
    urls_html = ["https://www.bis.org/ifc/publications.htm"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows = []
    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            for path, doc in data.get("list", {}).items():
                titulo = html.unescape(doc.get("short_title", ""))
                if not titulo: continue
                link = "https://www.bis.org" + doc.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"): link += ".htm"
                try: parsed_date = parser.parse(doc.get("publication_start_date", ""))
                except: continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "BPI"})
        except: continue
    for url in urls_html:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            soup = BeautifulSoup(res.text, 'html.parser')
            content_div = soup.find('div', id='cmsContent')
            if not content_div: continue
            for p in content_div.find_all('p'):
                a_tag = p.find('a')
                if not a_tag: continue
                titulo = a_tag.get_text(strip=True)
                href = a_tag.get('href', '')
                if not href or 'index.htm' in href: continue 
                link = "https://www.bis.org" + href if href.startswith('/') else href
                parsed_date = None
                try: parsed_date = parser.parse(p.get_text(strip=True).replace(titulo, '').strip(', '))
                except: pass
                if not parsed_date:
                    match = re.search(r'\b(20\d{2})\b', titulo)
                    if match: parsed_date = datetime.datetime(int(match.group(1)), 1, 1)
                if parsed_date and parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "BPI"})
        except: continue
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

# --- SECCIÓN: PUBLICACIONES INSTITUCIONALES ---
@st.cache_data(show_spinner=False)
def load_pub_inst_cef(start_date_str, end_date_str):
    url = "https://www.fsb.org/publications/key-regular-publications/"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows = []
    try:
        res = requests.get(url, headers=headers, timeout=15)
        soup = BeautifulSoup(res.text, 'html.parser')
        for section in soup.find_all('div', class_='wp-bootstrap-blocks-row'):
            h2 = section.find('h2')
            if not h2: continue
            base_title = h2.get_text(strip=True)
            # Latest
            latest_btn = section.find('button', class_='btn-primary')
            if latest_btn and latest_btn.find('a'):
                a_tag = latest_btn.find('a')
                link = "https://www.fsb.org" + a_tag['href'] if a_tag['href'].startswith('/') else a_tag['href']
                date_match = re.search(r'\((.*?)\)', a_tag.get_text())
                parsed_date = parser.parse(date_match.group(1)) if date_match else None
                if parsed_date and parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": f"{base_title}: Latest Report", "Link": link, "Organismo": "CEF"})
            # Previous
            dropdown = section.find('div', class_='dropdown-menu')
            if dropdown:
                for l in dropdown.find_all('a'):
                    year_text = l.get_text(strip=True)
                    try: parsed_date = datetime.datetime(int(year_text), 1, 1)
                    except: parsed_date = None
                    if parsed_date and parsed_date >= start_date:
                        rows.append({"Date": parsed_date, "Title": f"{base_title} ({year_text})", "Link": l['href'], "Organismo": "CEF"})
    except: pass
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_pub_inst_bpi(start_date_str, end_date_str):
    urls_api = ["https://www.bis.org/api/document_lists/annualeconomicreports.json", "https://www.bis.org/api/document_lists/quarterlyreviews.json"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows = []
    for url in urls_api:
        try:
            res = requests.get(url, headers=headers, timeout=15)
            data = res.json()
            for path, doc in data.get("list", {}).items():
                titulo = html.unescape(doc.get("short_title", ""))
                link = "https://www.bis.org" + doc.get("path", "")
                if not link.endswith(".htm") and not link.endswith(".pdf"): link += ".htm"
                try: parsed_date = parser.parse(doc.get("publication_start_date", ""))
                except: continue
                if parsed_date >= start_date:
                    rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "BPI"})
        except: continue
    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_country_reports_fmi(start_date_str, end_date_str):
    """Extractor FMI - Country Reports (Conexión Directa a Coveo API)"""
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    # 1. EL ENDPOINT Y LA LLAVE MAESTRA QUE DESCUBRISTE
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"
    
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }
    
    # 2. EL PAYLOAD (Falsificamos la petición del buscador)
    payload = {
        "aq": "@imfseries==\"IMF Staff Country Reports\"", # Filtro estricto por la Serie
        "numberOfResults": 100, # Cantidad a traer (Suficiente para un mes)
        "sortCriteria": "@imfdate descending" # Los más recientes primero
    }
    
    try:
        # Hacemos un POST directo a la base de datos de Coveo
        res = requests.post(url, headers=headers, json=payload, timeout=15)
        
        if res.status_code == 200:
            data = res.json()
            
            # 3. EXTRACCIÓN (Limpia y sin HTML)
            for item in data.get("results", []):
                titulo = item.get("title", "")
                link = item.get("clickUri", "")
                
                # La fecha viene en timestamp (milisegundos). Lo dividimos entre 1000 para segundos.
                raw_date = item.get("raw", {}).get("date")
                parsed_date = None
                if raw_date:
                    try:
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except: pass
                
                if not titulo or not link or not parsed_date: continue
                
                # Validamos contra la fecha del filtro de la app
                if parsed_date >= start_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "FMI"})
    except Exception as e:
        pass
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_press_releases_fmi(start_date_str, end_date_str):
    """Extractor FMI - Press Releases (Historial completo vía Coveo API)"""
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    # 1. El Endpoint y la llave que tú mismo descubriste
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"
    
    # 2. Inyección de Headers para evadir el bloqueo CORS
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",   # <--- LA LLAVE PARA ENTRAR
        "Referer": "https://www.imf.org/", # <--- CONFIRMA QUE "VENIMOS" DEL FMI
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }
    
    # 3. Payload: Agregamos el filtro estricto de idioma
    payload = {
        # Le pedimos PRs Y que el idioma sea inglés
        "aq": "@imftype==\"Press Release\" AND @syslanguage==\"English\"", 
        "numberOfResults": 150, 
        "sortCriteria": "@imfdate descending"
    }
    
    try:
        res = requests.post(url, headers=headers, json=payload, timeout=15)
        
        if res.status_code == 200:
            data = res.json()
            
            for item in data.get("results", []):
                titulo = item.get("title", "")
                link = item.get("clickUri", "")
                
                # Coveo entrega la fecha en formato Unix (Milisegundos). 
                # ¡Es perfecto porque no falla la conversión!
                raw_date = item.get("raw", {}).get("date")
                parsed_date = None
                if raw_date:
                    try:
                        # Convertimos de milisegundos a fecha normal
                        parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except: pass
                
                if not titulo or not link or not parsed_date: continue
                
                # Filtro final de fechas
                if parsed_date >= start_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "FMI"})
    except Exception as e:
        pass
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_country_reports_elibrary(start_date_str, end_date_str):
    """Extractor FMI - Country Reports (Bypass de Tapestry 5 AJAX Lazy-Loading)"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
    }
    
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    base_domain = "https://www.elibrary.imf.org"
    url_overview = f"{base_domain}/view/journals/002/002-overview.xml"
    
    try:
        # FASE 1: Extraer los tokens dinámicos de AJAX para los años recientes
        res = requests.get(url_overview, headers=headers, timeout=15)
        if res.status_code != 200: return pd.DataFrame()
        
        soup = BeautifulSoup(res.text, 'html.parser')
        
        ajax_links = []
        current_year = datetime.datetime.now().year
        # Buscamos los enlaces de expansión para el año actual y el anterior
        target_years = [str(current_year), str(current_year - 1)] 
        
        for li in soup.find_all('div', attrs={'data-toc-role': 'li'}):
            label_div = li.find('div', class_='label')
            if not label_div: continue
            
            texto_label = label_div.get_text()
            if any(year in texto_label for year in target_years):
                a_tag = li.find('a', class_='ajax-control')
                if a_tag and a_tag.has_attr('href'):
                    ajax_links.append(base_domain + a_tag['href'])
        
        # FASE 2: Interceptar y "deshidratar" las respuestas AJAX de Tapestry
        headers_ajax = headers.copy()
        headers_ajax['X-Requested-With'] = 'XMLHttpRequest' # Engañamos al framework
        headers_ajax['Accept'] = 'application/json, text/javascript, */*; q=0.01'
        
        for ajax_url in ajax_links:
            try:
                res_ajax = requests.get(ajax_url, headers=headers_ajax, timeout=15)
                if res_ajax.status_code != 200: continue
                
                data = res_ajax.json()
                
                # Extraemos el HTML inyectado dentro del nodo "zones"
                html_fragment = ""
                if "zones" in data:
                    for zone_id, html_content in data["zones"].items():
                        html_fragment += html_content
                        
                if not html_fragment: continue
                
                # FASE 3: Parsear el HTML revelado
                soup_fragment = BeautifulSoup(html_fragment, 'html.parser')
                
                for a_tag in soup_fragment.find_all('a', href=True):
                    href = a_tag['href']
                    titulo = a_tag.get_text(strip=True)
                    
                    # Filtro de sanidad: debe ser un artículo real
                    if '/view/journals/002/' in href and len(titulo) > 15:
                        link_real = base_domain + href if href.startswith('/') else href
                        
                        # Buscamos la fecha subiendo hasta 3 niveles en el DOM
                        date_str = ""
                        for padre in a_tag.find_parents(['div', 'li'], limit=3):
                            texto_padre = padre.get_text(separator=" ", strip=True)
                            
                            # Caza fechas en formatos "Mar 05, 2026" o "05 March 2026"
                            match = re.search(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2}?,?\s*\d{4}', texto_padre)
                            if not match:
                                match = re.search(r'\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}', texto_padre)
                                
                            if match:
                                date_str = match.group(0)
                                break # Encontramos la fecha, salimos del bucle
                                
                        parsed_date = None
                        if date_str:
                            try:
                                parsed_date = parser.parse(date_str)
                                if parsed_date.tzinfo is not None: parsed_date = parsed_date.replace(tzinfo=None)
                            except: pass
                            
                        # Evaluación final
                        if parsed_date and parsed_date >= start_date:
                            if not any(r['Link'] == link_real for r in rows):
                                rows.append({"Date": parsed_date, "Title": titulo, "Link": link_real, "Organismo": "FMI"})
            except:
                continue # Aislamiento de fallos
                
    except Exception as e:
        pass
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_pub_inst_fmi(start_date_str, end_date_str):
    """Extractor FMI - Vía directa por API Next.js (El Regalo)"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'application/json, text/plain, */*'
    }
    
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    # 1. CAZADOR DE BUILD ID (Para que tu código no caduque nunca)
    build_id = "OPXKbpp2La91iW-gTVkBX" # Tu regalo como plan de respaldo
    try:
        res_html = requests.get("https://www.imf.org/en/publications", headers=headers, timeout=15)
        # Buscamos el código dinámico oculto en la página principal
        match = re.search(r'"buildId":"([^"]+)"', res_html.text)
        if match:
            build_id = match.group(1)
    except:
        pass

    # 2. CONSTRUCCIÓN DE LOS ENLACES JSON DIRECTOS
    endpoints_json = [
        f"https://www.imf.org/_next/data/{build_id}/en/publications/fm.json",
        f"https://www.imf.org/_next/data/{build_id}/en/publications/weo.json",
        f"https://www.imf.org/_next/data/{build_id}/en/publications/gfsr.json"
    ]
    
    for url in endpoints_json:
        try:
            # Ahora pedimos el JSON limpio, evadiendo el HTML
            res = requests.get(url, headers=headers, timeout=15)
            if res.status_code != 200: continue
            data = res.json()
            
            # Buscador recursivo dentro del JSON
            def extraer_issues(obj):
                if isinstance(obj, dict):
                    if "issuePage" in obj and isinstance(obj["issuePage"], dict) and "results" in obj["issuePage"]:
                        for r in obj["issuePage"]["results"]: yield r
                    for k, v in obj.items(): yield from extraer_issues(v)
                elif isinstance(obj, list):
                    for item in obj: yield from extraer_issues(item)

            for issue in extraer_issues(data):
                titulo = issue.get("title", {}).get("jsonValue", {}).get("value", "")
                link_raw = issue.get("url", {}).get("url", "") or issue.get("url", {}).get("path", "")
                if not titulo or not link_raw: continue
                
                link_real = link_raw if link_raw.startswith("http") else "https://www.imf.org" + link_raw
                
                d_str = issue.get("publicationDate", {}).get("jsonValue", {}).get("value", "")
                if d_str:
                    try:
                        parsed_date = parser.parse(d_str)
                        if parsed_date.tzinfo is not None: parsed_date = parsed_date.replace(tzinfo=None)
                        if parsed_date >= start_date and not any(r['Link'] == link_real for r in rows):
                            rows.append({"Date": parsed_date, "Title": titulo, "Link": link_real, "Organismo": "FMI"})
                    except: pass
        except:
            continue
            
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_pub_inst_bm(start_date_str, end_date_str):
    """Extractor para Publicaciones Institucionales (Colecciones Específicas) del BM"""
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}
    
    # IDs exactos de las 3 colecciones
    scopes = [
        '4c48a649-7773-4d0f-b441-f5fc7e8d67f8', # Business Ready
        '09c5e8fc-187f-5c2f-a077-3e03044c7b62', # Perspectivas económicas mundiales
        '3d9bbbf6-c007-5043-b655-04d8a1cfbfb2'  # Tercera colección
    ]
    
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows = []
    
    # Iteramos sobre cada una de las colecciones
    for scope in scopes:
        page = 0
        while True:
            try:
                # Al pasarle el 'scope', la API restringe la búsqueda SOLO a esa colección
                params = {
                    'scope': scope,
                    'sort': 'dc.date.issued,DESC', 
                    'page': page, 
                    'size': 20
                }
                res = requests.get(base_url, headers=headers, params=params, timeout=15)
                data = res.json()
                
                objects = data.get('_embedded', {}).get('searchResult', {}).get('_embedded', {}).get('objects', [])
                if not objects: break
                
                items_found = 0
                for obj in objects:
                    item = obj.get('_embedded', {}).get('indexableObject', {})
                    meta = item.get('metadata', {})
                    
                    title = meta.get('dc.title', [{'value': ''}])[0].get('value', '')
                    date_s = meta.get('dc.date.issued', [{'value': ''}])[0].get('value', '')
                    
                    parsed_date = None
                    if date_s:
                        try: parsed_date = parser.parse(date_s)
                        except: pass
                    
                    if not parsed_date or parsed_date < start_date: continue
                    
                    link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get('value', '')
                    if not link: link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"
                    
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": title, "Link": link, "Organismo": "BM"})
                        items_found += 1
                
                if items_found == 0: break
                page += 1
                if page > 3: break # Límite de seguridad
                time.sleep(0.2)
            except:
                break
                
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df
    
    # --- SECCIÓN: INVESTIGACIÓN ---
    # BID (Working Papers en inglés)
@st.cache_data(show_spinner=False)
def load_investigacion_bid_en(start_date_str, end_date_str):
    """
    Extrae Working Papers del BID en inglés
    URL: https://publications.iadb.org/en?f%5B0%5D=type%3AWorking%20Papers
    """
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import datetime
    import pandas as pd
    import time
    import re
    from dateutil import parser

    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    
    # Configuración de paginación
    page = 0
    max_pages = 5  # Límite de páginas a extraer
    hay_resultados = True
    
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
    chrome_options.add_argument("--disable-blink-features=AutomationControlled")
    chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
    chrome_options.add_experimental_option('useAutomationExtension', False)

    try:
        print("🔍 Iniciando Selenium para BID Working Papers (EN)...")
        driver = webdriver.Chrome(options=chrome_options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        while page < max_pages and hay_resultados:
            # URL para Working Papers en inglés
            url = f"https://publications.iadb.org/en?f%5B0%5D=type%3AWorking%20Papers&page={page}"
            
            print(f"📄 Accediendo a página {page+1}: {url}")
            driver.get(url)

            try:
                WebDriverWait(driver, 20).until_not(
                    EC.title_contains("Just a moment")
                )
                print(f"✅ Página {page+1} cargada correctamente.")
            except:
                print(f"⚠️ La página {page+1} sigue mostrando 'Just a moment...', esperando...")
                time.sleep(10)

            time.sleep(5)
            html = driver.page_source
            soup = BeautifulSoup(html, 'html.parser')

            # Guardar HTML para depuración (solo primera página)
            if page == 0:
                with open("bid_debug_en.html", "w", encoding="utf-8") as f:
                    f.write(html)
                print("💾 HTML guardado en bid_debug_en.html")

            # Estrategias de búsqueda
            items = soup.find_all('div', class_='views-row')
            print(f"📚 Página {page+1} - Elementos encontrados: {len(items)}")

            if len(items) == 0:
                print(f"📭 No hay más elementos en página {page+1}")
                hay_resultados = False
                break

            # Mapeo de meses en inglés
            meses_en = {
                'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
                'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
            }

            for item in items:
                # ESTRATEGIA 1 (PRIORITARIA): Buscar específicamente el div con clase 'views-field-field-title'
                # Esta es la estructura exacta que vimos en el HTML
                title_elem = None
                title_container = item.find('div', class_='views-field-field-title')
                if title_container:
                    span_field = title_container.find('span', class_='field-content')
                    if span_field:
                        a_tag = span_field.find('a')
                        if a_tag:
                            title_elem = a_tag
                            print(f"  ✅ Título encontrado con estrategia 1")

                # ESTRATEGIA 2: Buscar span.field-content > a (estructura genérica)
                if not title_elem:
                    span_field = item.find('span', class_='field-content')
                    if span_field:
                        a_tag = span_field.find('a')
                        if a_tag and a_tag.get_text(strip=True):
                            title_elem = a_tag
                            print(f"  ✅ Título encontrado con estrategia 2")

                # ESTRATEGIA 3: Buscar cualquier enlace con texto largo
                if not title_elem:
                    for a_tag in item.find_all('a', href=True):
                        texto = a_tag.get_text(strip=True)
                        if len(texto) > 30:
                            title_elem = a_tag
                            print(f"  ✅ Título encontrado con estrategia 3")
                            break

                if not title_elem:
                    print(f"  ⚠️ No se encontró título en elemento")
                    continue

                titulo = title_elem.get_text(strip=True)
                link = title_elem['href']
                if not link.startswith('http'):
                    link = "https://publications.iadb.org" + link

                print(f"  📌 Título extraído: '{titulo[:100]}...'")

                # Extraer fecha - VERSIÓN MEJORADA
                parsed_date = None
                
                # Buscar específicamente el contenedor de fecha
                date_container = item.find('div', class_='views-field-field-date-issued-text')
                if date_container:
                    date_span = date_container.find('span', class_='field-content')
                    if date_span:
                        date_text = date_span.get_text(strip=True)
                        print(f"  📅 Texto de fecha (específico): {date_text}")
                        
                        # Intentar parsear con regex (ej: "Mar 2026")
                        match = re.search(r'([A-Za-z]{3,9})\s+(\d{4})', date_text)
                        if match:
                            mes_str, año_str = match.groups()
                            mes_num = meses_en.get(mes_str.lower()[:3])
                            if mes_num:
                                parsed_date = datetime.datetime(int(año_str), mes_num, 1)
                                print(f"  ✅ Fecha parseada: {parsed_date}")
                
                # Fallback: buscar cualquier span con texto de fecha
                if not parsed_date:
                    for span in item.find_all('span'):
                        text = span.get_text(strip=True)
                        match = re.search(r'([A-Za-z]{3,9})\s+(\d{4})', text)
                        if match:
                            mes_str, año_str = match.groups()
                            mes_num = meses_en.get(mes_str.lower()[:3])
                            if mes_num:
                                parsed_date = datetime.datetime(int(año_str), mes_num, 1)
                                print(f"  ✅ Fecha parseada (fallback): {parsed_date}")
                                break

                if not parsed_date:
                    print(f"  ⚠️ No se pudo extraer fecha")
                    continue

                # Filtrar por fecha
                if parsed_date < start_date or parsed_date > end_date:
                    continue

                # Evitar duplicados
                if not any(r['Link'] == link for r in rows):
                    rows.append({
                        "Date": parsed_date,
                        "Title": titulo,
                        "Link": link,
                        "Organismo": "BID (Inglés)"
                    })
                    print(f"  ✅ Documento AGREGADO: {titulo[:50]}...")

            page += 1
            print(f"➡️ Avanzando a página {page+1}...\n")

        driver.quit()

    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.drop_duplicates(subset=['Link'])
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
        print(f"\n✅ Documentos BID (EN) encontrados en {page} páginas: {len(df)}")
    else:
        print("\n⚠️ No se encontraron documentos del BID (EN)")

    return df
@st.cache_data(show_spinner=False)
def load_investigacion_bm(start_date_str, end_date_str):
    """Extractor para Investigación del BM (Filtra y excluye los que son 'Reports')"""
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}
    
    # ID exacto de la comunidad de Investigación
    scope_id = '06251f8a-62c2-59fb-add5-ec0993fc20d9'
    
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    
    rows, page = [], 0
    while True:
        try:
            params = {
                'scope': scope_id, 
                'sort': 'dc.date.issued,DESC', 
                'page': page, 
                'size': 20
            }
            res = requests.get(base_url, headers=headers, params=params, timeout=15)
            data = res.json()
            
            objects = data.get('_embedded', {}).get('searchResult', {}).get('_embedded', {}).get('objects', [])
            if not objects: break
            
            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})
                
                # Extraer Título y Fecha
                title = meta.get('dc.title', [{'value': ''}])[0].get('value', '')
                date_s = meta.get('dc.date.issued', [{'value': ''}])[0].get('value', '')
                
                parsed_date = None
                if date_s:
                    try: parsed_date = parser.parse(date_s)
                    except: pass
                
                if not parsed_date or parsed_date < start_date: continue
                
                # --- NUEVO FILTRO ANTI-REPORTES ---
                # Buscamos en el abstract o en la descripción general
                abstract_list = meta.get('dc.description.abstract', [])
                desc_list = meta.get('dc.description', [])
                
                description = ""
                if abstract_list: description = abstract_list[0].get('value', '').lower()
                elif desc_list: description = desc_list[0].get('value', '').lower()
                
                # Si la palabra exacta "report" está en la descripción, lo saltamos
                # Usamos \b para que sea la palabra exacta y no algo como "reporting"
                if re.search(r'\breport\b', description):
                    continue
                # ----------------------------------
                
                # Link permanente
                link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get('value', '')
                if not link: link = f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"
                
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": title, "Link": link, "Organismo": "BM"})
                    items_found += 1
            
            if items_found == 0: break
            page += 1
            if page > 3: break # Límite para evitar búsquedas infinitas
            time.sleep(0.2)
        except:
            break
            
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        if df["Date"].dt.tz is not None: df["Date"] = df["Date"].dt.tz_convert(None)
        df = df.sort_values("Date", ascending=False)
    return df
# --- SECCIÓN: DISCURSOS ---
@st.cache_data(show_spinner=False)
@st.cache_data(show_spinner=False)
def load_discursos_fmi(start_date_str, end_date_str):
    """Extractor FMI - Discursos (Vía Coveo API) con Limpieza Avanzada"""
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
        end_date = datetime.datetime.now()

    rows = []
    url = "https://imfproduction561s308u.org.coveo.com/rest/search/v2?organizationId=imfproduction561s308u"
    headers = {
        "Authorization": "Bearer xx742a6c66-f427-4f5a-ae1e-770dc7264e8a",
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Origin": "https://www.imf.org",
        "Referer": "https://www.imf.org/",
        "User-Agent": "Mozilla/5.0"
    }
    payload = {"aq": "@imftype==\"Speech\" AND @syslanguage==\"English\"", "numberOfResults": 150, "sortCriteria": "@imfdate descending"}

    try:
        res = requests.post(url, headers=headers, json=payload, timeout=15)
        if res.status_code == 200:
            data = res.json()
            for item in data.get("results", []):
                titulo_raw = item.get("title", "").strip()
                link = item.get("clickUri", "")
                raw_data = item.get("raw", {})
                raw_date = raw_data.get("date")
                
                # Extraemos el autor
                autor = raw_data.get("imfspeaker", "")
                if isinstance(autor, list) and len(autor) > 0: autor = autor[0]
                autor = clean_author_name(autor)
                
                parsed_date = None
                if raw_date:
                    try: parsed_date = datetime.datetime.fromtimestamp(raw_date / 1000.0)
                    except: pass
                if not titulo_raw or not link or not parsed_date: continue
                
                # ---------------------------------------------------------
                # LIMPIEZA NIVEL DIOS (Adiós comillas y subtítulos largos)
                # ---------------------------------------------------------
                titulo_limpio = titulo_raw
                
                # 1. Quitar la "cola" institucional (ej: " - Keynote Speech by...")
                patron_sufijo = re.compile(r"(?i)\s*[\-–—]\s*.*?(speech|remarks|statement|address)\s+by\s+.*$")
                titulo_limpio = patron_sufijo.sub("", titulo_limpio).strip()
                
                # 2. Quitar comillas sobrantes que hayan quedado expuestas
                titulo_limpio = titulo_limpio.strip('"').strip("'").strip()
                
                # 3. Formatear "Autor: Título" evitando duplicados
                if autor:
                    patron_inicio = re.compile(rf"^{re.escape(autor)}\s*[:\-]\s*", re.IGNORECASE)
                    titulo_limpio = patron_inicio.sub("", titulo_limpio).strip('"').strip("'").strip()
                    titulo_final = f"{autor}: {titulo_limpio}"
                else:
                    titulo_final = titulo_limpio
                # ---------------------------------------------------------

                if start_date <= parsed_date <= end_date:
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": titulo_final, "Link": link, "Organismo": "FMI"})
    except: pass
    
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_ecb(start_date_str, end_date_str):
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    try: 
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
        end_date = datetime.datetime.strptime(end_date_str, '%d.%m.%Y')
        anios_num = list(range(start_date.year, end_date.year + 1))
    except: anios_num = [2026, 2025, 2024]
    for year in anios_num:
        url = f"https://www.ecb.europa.eu/press/key/date/{year}/html/index.en.html"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            for a in soup.find_all('a', href=True):
                href = a['href']
                if f'/press/key/date/{year}/html/' in href and href.endswith('.html') and 'index' not in href:
                    link = "https://www.ecb.europa.eu" + href if href.startswith('/') else href
                    titulo_raw = a.get_text(strip=True)
                    if len(titulo_raw) < 5: continue
                    parent = a.find_parent(['dd', 'div', 'li'])
                    if not parent: continue
                    dt = parent.find_previous_sibling('dt')
                    fecha_str = dt.get_text(strip=True) if dt else ""
                    try: parsed_date = parser.parse(fecha_str)
                    except: continue
                    autor = ""
                    sub = parent.find('div', class_='subtitle')
                    if sub:
                        match = re.search(r'\b(?:by|with)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', sub.get_text(separator=' ', strip=True))
                        if match: autor = clean_author_name(match.group(1))
                    final_t = f"{autor}: {titulo_raw}" if autor and autor not in titulo_raw else titulo_raw
                    if not any(r['Link'] == link for r in rows):
                        rows.append({"Date": parsed_date, "Title": final_t, "Link": link, "Organismo": "ECB (Europa)"})
        except: pass
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_bis():
    urls = ["https://www.bis.org/api/document_lists/cbspeeches.json", "https://www.bis.org/api/document_lists/bcbs_speeches.json", "https://www.bis.org/api/document_lists/mgmtspeeches.json"]
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    for url in urls:
        try:
            response = requests.get(url, headers=headers, timeout=10)
            data = response.json()
            for path, speech in data.get("list", {}).items():
                title = html.unescape(speech.get("short_title", ""))
                date_str = speech.get("publication_start_date", "")
                link = "https://www.bis.org" + path + (".htm" if not path.endswith(".htm") else "")
                rows.append({"Date": date_str, "Title": title, "Link": link, "Organismo": "BPI"})
        except: continue
    df = pd.DataFrame(rows).drop_duplicates(subset=['Link']) if rows else pd.DataFrame()
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_bbk(start_date_str, end_date_str):
    base_url = "https://www.bundesbank.de/action/en/730564/bbksearch"
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows, page = [], 0
    while True:
        params = {'sort': 'bbksortdate desc', 'dateFrom': start_date_str, 'dateTo': end_date_str, 'pageNumString': str(page)}
        try: response = requests.get(base_url, headers=headers, params=params, timeout=10)
        except: break 
        soup = BeautifulSoup(response.text, 'html.parser')
        items = soup.find_all('li', class_='resultlist__item')
        if not items: break 
        for item in items:
            fecha_tag = item.find('span', class_='metadata__date')
            fecha_str = fecha_tag.text.strip() if fecha_tag else ""
            author_tag = item.find('span', class_='metadata__authors')
            author_str = clean_author_name(author_tag.text) if author_tag else ""
            data_div = item.find('div', class_='teasable__data')
            link, titulo = "", ""
            if data_div and data_div.find('a'):
                a_tag = data_div.find('a')
                link = "https://www.bundesbank.de" + a_tag.get('href', '') if a_tag.get('href', '').startswith('/') else a_tag.get('href', '')
                if a_tag.find('span', class_='link__label'): titulo = a_tag.find('span', class_='link__label').text.strip()
            if author_str and author_str not in titulo: titulo = f"{author_str}: {titulo}"
            if fecha_str and titulo: rows.append({"Date": fecha_str, "Title": titulo, "Link": link, "Organismo": "BBk (Alemania)"})
        if len(items) < 10: break
        page += 1
        time.sleep(0.3) 
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"], format='%d.%m.%Y', errors='coerce')
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_pboc(start_date_str, end_date_str):
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        url = "https://www.pbc.gov.cn/en/3688110/3688175/index.html" if page == 1 else f"https://www.pbc.gov.cn/en/3688110/3688175/0180081b-{page}.html"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            res.encoding = 'utf-8' 
            soup = BeautifulSoup(res.text, 'html.parser')
            items = soup.find_all('div', class_='ListR')
            if not items: break
            items_found = 0
            for item in items:
                date_span = item.find('span', class_='prhhdata')
                a_tag = item.find('a')
                if not date_span or not a_tag: continue
                try: parsed_date = parser.parse(date_span.get_text(strip=True))
                except: continue
                titulo_raw = html.unescape(a_tag.get('title', a_tag.get_text(strip=True)))
                link = "https://www.pbc.gov.cn" + a_tag.get('href', '') if a_tag.get('href', '').startswith('/') else a_tag.get('href', '')
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "PBoC (China)"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.5) 
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_fed(anios_num):
    headers = {'User-Agent': 'Mozilla/5.0'}
    rows = []
    for year in anios_num:
        url = f"https://www.federalreserve.gov/newsevents/{year}-speeches.htm"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            if res.status_code == 404:
                url = "https://www.federalreserve.gov/newsevents/speeches.htm"
                res = requests.get(url, headers=headers, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            for a_tag in soup.find_all('a', href=True):
                if '/newsevents/speech/' in a_tag['href']:
                    link = "https://www.federalreserve.gov" + a_tag['href'] if a_tag['href'].startswith('/') else a_tag['href']
                    titulo = a_tag.get_text(strip=True)
                    parent = a_tag.find_parent('div', class_='row') or a_tag.parent
                    text = parent.get_text(separator=' | ', strip=True)
                    date_m = re.search(r'(\d{1,2}/\d{1,2}/\d{4}|\w+\s\d{1,2},\s\d{4})', text)
                    if date_m:
                        try:
                            parsed_date = parser.parse(date_m.group(1))
                            if parsed_date.year not in anios_num: continue
                            rows.append({"Date": parsed_date, "Title": titulo, "Link": link, "Organismo": "Fed (Estados Unidos)"})
                        except: pass
        except: pass
    df = pd.DataFrame(rows).drop_duplicates(subset=['Link']) if rows else pd.DataFrame()
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_bdf(start_date_str, end_date_str):
    base_url = "https://www.banque-france.fr/en/governor-interventions"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 0
    while True:
        try:
            response = requests.get(base_url, headers=headers, params={'category[7052]': '7052', 'page': page}, timeout=12)
            soup = BeautifulSoup(response.text, 'html.parser')
            cards = soup.find_all('div', class_=lambda c: c and 'card' in c)
            if not cards: break
            items_found = 0
            for card in cards:
                a = card.find('a', href=True)
                if not a or not a.find('span', class_='title-truncation'): continue
                titulo_raw, link = a.find('span', class_='title-truncation').get_text(strip=True), "https://www.banque-france.fr" + a['href']
                date_s = card.find('small')
                if not date_s: continue
                fecha_clean = re.sub(r'(\d+)(st|nd|rd|th)\s+of\s+', r'\1 ', date_s.get_text(strip=True))
                try: parsed_date = parser.parse(fecha_clean)
                except: continue
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "BdF (Francia)"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.3)
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_bm(start_date_str, end_date_str):
    base_url = "https://openknowledge.worldbank.org/server/api/discover/search/objects"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 0
    while True:
        try:
            res = requests.get(base_url, headers=headers, params={'scope': 'b6a50016-276d-56d3-bbe5-891c8d18db24', 'sort': 'dc.date.issued,DESC', 'page': page, 'size': 20}, timeout=12)
            objects = res.json().get('_embedded', {}).get('searchResult', {}).get('_embedded', {}).get('objects', [])
            if not objects: break
            items_found = 0
            for obj in objects:
                item = obj.get('_embedded', {}).get('indexableObject', {})
                meta = item.get('metadata', {})
                title = meta.get('dc.title', [{'value': ''}])[0].get('value', '')
                date_s = meta.get('dc.date.issued', [{'value': ''}])[0].get('value', '')
                try: parsed_date = parser.parse(date_s)
                except: continue
                link = meta.get('dc.identifier.uri', [{'value': ''}])[0].get('value', '') or f"https://openknowledge.worldbank.org/entities/publication/{item.get('id', '')}"
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": title, "Link": link, "Organismo": "BM"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.3)
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_boc(start_date_str, end_date_str):
    base_url = "https://www.bankofcanada.ca/press/speeches/"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        try:
            res = requests.get(base_url, headers=headers, params={'mt_page': page}, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            articles = soup.find_all('div', class_=lambda c: c and ('mtt-result' in c or 'media' in c))
            if not articles: break
            items_found = 0
            for art in articles:
                h3 = art.find('h3', class_='media-heading')
                if not h3 or not h3.find('a'): continue
                titulo_raw, link = h3.find('a').text.strip(), h3.find('a')['href']
                date_s = art.find('span', class_='media-date')
                try: parsed_date = parser.parse(date_s.text.strip())
                except: continue
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "BoC (Canadá)"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.3)
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_boj(start_date_str, end_date_str):
    base_url = "https://www.boj.or.jp/en/about/press/index.htm"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows = []
    try:
        response = requests.get(base_url, headers=headers, timeout=12)
        soup = BeautifulSoup(response.text, 'html.parser')
        table = soup.find('table', class_='js-tbl')
        if table:
            for tr in table.find('tbody').find_all('tr'):
                tds = tr.find_all('td')
                if len(tds) < 3: continue
                try: parsed_date = parser.parse(tds[0].get_text(strip=True).replace('\xa0', ' '))
                except: continue
                if parsed_date < start_date: continue
                a_tag = tds[2].find('a', href=True)
                if not a_tag: continue
                titulo_raw = a_tag.get_text(strip=True).strip('"')
                link = "https://www.boj.or.jp" + a_tag['href'] if a_tag['href'].startswith('/') else a_tag['href']
                rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "BoJ (Japón)"})
    except: pass
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner=False)
def load_data_cef(start_date_str, end_date_str):
    base_url = "https://www.fsb.org/press/speeches-and-statements/"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try: start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except: start_date = datetime.datetime(2000, 1, 1)
    rows, page = [], 1
    while True:
        url = f"{base_url}?dps_paged={page}"
        try:
            res = requests.get(url, headers=headers, timeout=12)
            soup = BeautifulSoup(res.text, 'html.parser')
            items = soup.find_all('div', class_='post-excerpt')
            if not items: break
            items_found = 0
            for item in items:
                title_tag = item.find('div', class_='post-title')
                if not title_tag or not title_tag.find('a'): continue
                a = title_tag.find('a')
                titulo_raw, link = a.get_text(strip=True), a['href']
                date_tag = item.find('div', class_='post-date')
                try: parsed_date = parser.parse(date_tag.get_text(strip=True))
                except: continue
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": parsed_date, "Title": titulo_raw, "Link": link, "Organismo": "CEF"})
                    items_found += 1
            if items_found == 0 or (rows and rows[-1]['Date'] < start_date): break
            page += 1
            time.sleep(0.3)
        except: break
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

# ==========================================
# EXPORTACIÓN A WORD
# ==========================================
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    c = docx.oxml.shared.OxmlElement('w:color'); c.set(docx.oxml.shared.qn('w:val'), '0000EE'); rPr.append(c)
    u = docx.oxml.shared.OxmlElement('w:u'); u.set(docx.oxml.shared.qn('w:val'), 'single'); rPr.append(u)
    b = docx.oxml.shared.OxmlElement('w:b'); rPr.append(b)
    
    for s in ['w:sz', 'w:szCs']:
        sz = docx.oxml.shared.OxmlElement(s); sz.set(docx.oxml.shared.qn('w:val'), '28'); rPr.append(sz)
        
    rFonts = docx.oxml.shared.OxmlElement('w:rFonts'); rFonts.set(docx.oxml.shared.qn('w:ascii'), 'Calibri'); rFonts.set(docx.oxml.shared.qn('w:hAnsi'), 'Calibri'); rPr.append(rFonts)
    t = docx.oxml.shared.OxmlElement('w:t'); t.text = text; new_run.append(rPr); new_run.append(t); hyperlink.append(new_run); paragraph._p.append(hyperlink)

def generate_word(df, title="Boletín Mensual", subtitle=""):
    doc = Document()
    h = doc.add_heading(title, 0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle); run.font.name, run.font.size = 'Calibri', Pt(14)
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=len(df.columns)-1)
    table.style = 'Table Grid'
    
    cols = [c for c in df.columns if c != 'Link']
    
    for idx, name in enumerate(cols):
        p = table.rows[0].cells[idx].paragraphs[0]
        run = p.add_run(name)
        run.font.name = 'Calibri'
        run.font.size = Pt(14) 
        run.bold = True
        
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            p = cells[i].paragraphs[0]
            if col == 'Nombre de Documento': 
                add_hyperlink(p, str(row[col]), str(row['Link']))
            else:
                run = p.add_run(str(row[col]))
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
                run.bold = True

    if 'Tipo de Documento' in df.columns and 'Organismo' in df.columns:
        col_tipo = cols.index('Tipo de Documento')
        col_org = cols.index('Organismo')
        
        start_row = 1
        while start_row <= len(df):
            cat_val = df.iloc[start_row - 1]['Tipo de Documento']
            org_val = df.iloc[start_row - 1]['Organismo']
            end_row = start_row
            
            if cat_val == "Discursos":
                table.cell(start_row, col_org).text = "" 
                while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == "Discursos":
                    table.cell(end_row + 1, col_org).text = "" 
                    end_row += 1
                
                if end_row > start_row:
                    target_cell = table.cell(start_row, col_org)
                    target_cell.merge(table.cell(end_row, col_org))
                
                start_row = end_row + 1
                continue
                
            while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == cat_val and df.iloc[end_row]['Organismo'] == org_val:
                table.cell(end_row + 1, col_org).text = "" 
                end_row += 1
                
            if end_row > start_row:
                target_cell = table.cell(start_row, col_org)
                target_cell.merge(table.cell(end_row, col_org))
                target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
                
            start_row = end_row + 1

        start_row = 1
        while start_row <= len(df):
            cat_val = df.iloc[start_row - 1]['Tipo de Documento']
            end_row = start_row
            
            while end_row < len(df) and df.iloc[end_row]['Tipo de Documento'] == cat_val:
                table.cell(end_row + 1, col_tipo).text = ""
                end_row += 1
            
            if end_row > start_row:
                target_cell = table.cell(start_row, col_tipo)
                target_cell.merge(table.cell(end_row, col_tipo))
                target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
                
            start_row = end_row + 1
                
    out = BytesIO(); doc.save(out); out.seek(0); return out

# ==========================================
# INTERFAZ DE USUARIO Y MAIN
# ==========================================
try: 
    st.sidebar.image("logo_banxico.png", use_container_width=True)
except: 
    st.sidebar.markdown("### 🏦 BANCO DE MÉXICO")

st.sidebar.markdown("---")
st.sidebar.header("Menú de Navegación")
modo_app = st.sidebar.radio("", ["Boletín", "Categorías"], key="menu_principal") 
st.sidebar.markdown("---")

anios_str = ["2026", "2025", "2024", "2023", "2022"]
meses_dict = {
    "Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6,
    "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12
}

# --- LISTAS DINÁMICAS DE ORGANISMOS ---
orgs_discursos = ["BBk (Alemania)", "BdE (España)", "BdF (Francia)", "BM", "BoC (Canadá)", "BoJ (Japón)", "BPI", "CEF", "ECB (Europa)", "Fed (Estados Unidos)", "FMI", "PBoC (China)"]
orgs_reportes = ["BID", "BM", "BPI", "CEF", "FEM", "OCDE"]
orgs_pub_inst = ["BM", "BPI", "CEF", "CEMLA", "FMI", "G20", "OCDE", "OEI"] 
orgs_investigacion = ["BID", "BM", "BPI", "CEMLA", "FMI", "OCDE"]

if modo_app == "Boletín":
    st.title("Generador de Boletín Mensual")
    st.markdown("Extrae y unifica documentos de todas las categorías y organismos por mes."); st.markdown("---")
    
    c1, c2 = st.columns(2)
    m_sel = c1.multiselect("Mes(es)", options=list(meses_dict.keys()))
    a_sel = c2.multiselect("Año(s)", options=anios_str, default=["2026"])
    
    if st.button("📄 Generar Boletín Mensual", type="primary"):
        if not m_sel or not a_sel: 
            st.warning("⚠️ Selecciona mes y año.")
        else:
            m_num = [meses_dict[m] for m in m_sel]
            a_num = [int(a) for a in a_sel]
            sd = f"01.{min(m_num):02d}.{min(a_num)}"
            ed = f"{calendar.monthrange(max(a_num), max(m_num))[1]:02d}.{max(m_num):02d}.{max(a_num)}"
            
            all_dfs = []
            prog = st.progress(0)
            txt = st.empty()
            
            total_pasos = len(orgs_discursos) + len(orgs_reportes) + len(orgs_pub_inst) + len(orgs_investigacion)
            paso_actual = 0
            
            # 1. BARRIDO DE DISCURSOS
            for org in orgs_discursos:
                txt.text(f"Procesando Discursos: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BPI": df = load_data_bis()
                    elif org == "ECB (Europa)": df = load_data_ecb(sd, ed)
                    elif org == "FMI": df = load_discursos_fmi(sd, ed)
                    elif org == "BBk (Alemania)": df = load_data_bbk(sd, ed)
                    elif org == "Fed (Estados Unidos)": df = load_data_fed(a_num)
                    elif org == "BdF (Francia)": df = load_data_bdf(sd, ed)
                    elif org == "BM": df = load_data_bm(sd, ed)
                    elif org == "BoC (Canadá)": df = load_data_boc(sd, ed)
                    elif org == "BoJ (Japón)": df = load_data_boj(sd, ed)
                    elif org == "CEF": df = load_data_cef(sd, ed)
                    elif org == "PBoC (China)": df = load_data_pboc(sd, ed)
                except Exception as e: pass 
                
                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty: 
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Discursos"
                        all_dfs.append(df_f)
                paso_actual += 1; prog.progress(paso_actual / total_pasos)

            # 2. BARRIDO DE REPORTES
            for org in orgs_reportes:
                txt.text(f"Procesando Reportes: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BID": df = load_reportes_bid_en(sd, ed)
                    elif org == "BM": df = load_reportes_bm(sd, ed) # <--- AGRÉGALO AQUÍ
                    elif org == "BPI": df = load_reportes_bpi(sd, ed)
                    elif org == "CEF": df = load_reportes_cef(sd, ed)
                    elif org == "OCDE": df = load_reportes_ocde(sd, ed)
                except Exception as e: pass
                
                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty: 
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Reportes"
                        all_dfs.append(df_f)
                paso_actual += 1; prog.progress(paso_actual / total_pasos)
                
            # 3. BARRIDO DE PUBLICACIONES INSTITUCIONALES 
            for org in orgs_pub_inst:
                txt.text(f"Procesando Pub. Institucionales: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BPI": df = load_pub_inst_bpi(sd, ed)
                    elif org == "CEF": df = load_pub_inst_cef(sd, ed)
                    elif org == "BM": df = load_pub_inst_bm(sd, ed)
                    elif org == "FMI": 
                        # 1. SSG - JSON Estático (WEO, Fiscal Monitor)
                        df_flagships = load_pub_inst_fmi(sd, ed)
                        
                        # 2. SSG - JSON Estático (Comunicados)
                        df_prs = load_press_releases_fmi(sd, ed)
                        
                        # 3. CSR API - Coveo (Country Reports)
                        df_crs = load_country_reports_fmi(sd, ed) # <-- LA NUEVA API
                        
                        # Unión
                        dfs_a_unir = [d for d in [df_flagships, df_prs, df_crs] if not d.empty]
                        if dfs_a_unir:
                            df = pd.concat(dfs_a_unir, ignore_index=True)
                            df = df.sort_values("Date", ascending=False)
                            
                except Exception as e: pass 
                
                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty: 
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Publicaciones Institucionales"
                        all_dfs.append(df_f)
                paso_actual += 1; prog.progress(paso_actual / total_pasos)

            # 4. BARRIDO DE INVESTIGACIÓN
            for org in orgs_investigacion:
                txt.text(f"Procesando Investigación: {org}...")
                df = pd.DataFrame()
                try:
                    if org == "BPI": df = load_investigacion_bpi(sd, ed)
                    elif org == "BM": df = load_investigacion_bm(sd, ed)
                except Exception as e: pass 
                
                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty: 
                        df_f['Organismo'] = org
                        df_f['Categoría'] = "Investigación"
                        all_dfs.append(df_f)
                paso_actual += 1; prog.progress(paso_actual / total_pasos)
            
            txt.empty()
            prog.empty()
            
            # --- CONSOLIDACIÓN FINAL ---
            if all_dfs:
                f_df = pd.concat(all_dfs, ignore_index=True)
                
                # --- PREPARACIÓN PARA EL WORD (Orden Institucional) ---
                df_rep = f_df[f_df['Categoría'] == "Reportes"].copy()
                df_pub = f_df[f_df['Categoría'] == "Publicaciones Institucionales"].copy()
                df_inv = f_df[f_df['Categoría'] == "Investigación"].copy()
                df_disc = f_df[f_df['Categoría'] == "Discursos"].copy()
                
                if not df_rep.empty: df_rep = df_rep.sort_values(by=["Organismo", "Title"], ascending=[True, True])
                if not df_pub.empty: df_pub = df_pub.sort_values(by=["Organismo", "Title"], ascending=[True, True])
                if not df_inv.empty: df_inv = df_inv.sort_values(by=["Organismo", "Title"], ascending=[True, True])
                if not df_disc.empty: df_disc = df_disc.sort_values(by=["Title"], ascending=[True])
                
                f_df_word = pd.concat([df_rep, df_pub, df_inv, df_disc], ignore_index=True)
                f_df_word = f_df_word[['Categoría', 'Organismo', 'Title', 'Link']]
                f_df_word = f_df_word.rename(columns={"Categoría": "Tipo de Documento", "Title": "Nombre de Documento"})
                
                st.success(f"Se consolidaron **{len(f_df)}** documentos en total.")
                word = generate_word(f_df_word, subtitle=", ".join(m_sel) + " " + ", ".join(a_sel))
                st.download_button("📄 Descargar Boletín", word, f"Boletin_{'_'.join(m_sel)}.docx")
                
                # --- PREPARACIÓN PARA LA VISTA PREVIA (Orden Cronológico + Columna Fecha) ---
                disp = f_df.copy()
                disp = disp.sort_values(by="Date", ascending=False) # Orden cronológico
                disp["Fecha"] = disp["Date"].dt.strftime('%d/%m/%Y') # Formatear fecha para que se vea bonita
                disp["Nombre de Documento"] = disp.apply(lambda x: f"[{x['Title']}]({x['Link']})", axis=1)
                disp = disp.rename(columns={"Categoría": "Tipo de Documento"})
                
                # Mostrar en pantalla incluyendo la columna 'Fecha'
                st.markdown(disp[["Fecha", "Tipo de Documento", "Organismo", "Nombre de Documento"]].to_markdown(index=False), unsafe_allow_html=True)
            else: 
                st.warning("No se encontraron documentos para los criterios seleccionados.")

elif modo_app == "Categorías":
    st.title("Documentos de Organismos Internacionales")
    tipo_doc = st.sidebar.selectbox("Tipo de Documento", ["Discursos", "Reportes", "Investigación", "Publicaciones Institucionales"])
    
    # Construcción segura de las listas de interfaz
    if tipo_doc == "Discursos": orgs_list = ["Todos"] + sorted(orgs_discursos)
    elif tipo_doc == "Reportes": orgs_list = ["Todos"] + sorted(orgs_reportes)
    elif tipo_doc == "Investigación": orgs_list = ["Todos"] + sorted(orgs_investigacion)
    elif tipo_doc == "Publicaciones Institucionales": orgs_list = ["Todos"] + sorted(orgs_pub_inst)
    else: orgs_list = ["Todos"] + sorted(list(set(orgs_discursos + orgs_reportes + orgs_investigacion + orgs_pub_inst)))
        
    organismo_seleccionado = st.sidebar.selectbox("Organismo", orgs_list)
    
    c1, c2 = st.columns(2)
    m_sel = c1.multiselect("Mes(es)", options=list(meses_dict.keys()))
    a_sel = c2.multiselect("Año(s)", options=anios_str, default=["2026"])
    
    if st.button("🔍 Buscar", type="primary"):
        if not m_sel or not a_sel:
            st.warning("⚠️ Selecciona mes y año.")
        else:
            m_num = [meses_dict[m] for m in m_sel]
            a_num = [int(a) for a in a_sel]
            sd = f"01.{min(m_num):02d}.{min(a_num)}"
            ed = f"{calendar.monthrange(max(a_num), max(m_num))[1]:02d}.{max(m_num):02d}.{max(a_num)}"
            
            target_orgs = orgs_list[1:] if organismo_seleccionado == "Todos" else [organismo_seleccionado]
            dfs_comb = []
            progreso = st.progress(0)
            txt = st.empty()
            
            for i, o in enumerate(target_orgs):
                txt.text(f"Extrayendo: {o}...")
                df = pd.DataFrame()
                try:
                    # --- LÓGICA DE EXTRACCIÓN POR TIPO ---
                    if tipo_doc == "Discursos":
                        if o == "BPI": df = load_data_bis()
                        elif o == "ECB (Europa)": df = load_data_ecb(sd, ed)
                        elif o == "BBk (Alemania)": df = load_data_bbk(sd, ed)
                        elif o == "Fed (Estados Unidos)": df = load_data_fed(a_num)
                        elif o == "BdF (Francia)": df = load_data_bdf(sd, ed)
                        elif o == "BM": df = load_data_bm(sd, ed)
                        elif o == "BoC (Canadá)": df = load_data_boc(sd, ed)
                        elif o == "BoJ (Japón)": df = load_data_boj(sd, ed)
                        elif o == "CEF": df = load_data_cef(sd, ed)
                        elif o == "FMI": df = load_discursos_fmi(sd, ed) # <--- AQUÍ ESTÁ EL FMI
                        elif o == "PBoC (China)": df = load_data_pboc(sd, ed)
                    
                    elif tipo_doc == "Reportes":
                        if o == "BID": 
                            dfs_bid = []
                            try: dfs_bid.append(load_reportes_bid(sd, ed))
                            except: pass
                            try: dfs_bid.append(load_reportes_bid_en(sd, ed))
                            except: pass
                            dfs_bid = [d for d in dfs_bid if not d.empty]
                            if dfs_bid: df = pd.concat(dfs_bid, ignore_index=True).drop_duplicates(subset=['Link'])
                        elif o == "BM": df = load_reportes_bm(sd, ed)
                        elif o == "BPI": df = load_reportes_bpi(sd, ed)
                        elif o == "CEF": df = load_reportes_cef(sd, ed)
                        elif o == "OCDE": df = load_reportes_ocde(sd, ed)
                        
                    elif tipo_doc == "Investigación":
                        if o == "BID": 
                            dfs_bid = []
                            try: dfs_bid.append(load_investigacion_bid(sd, ed))
                            except: pass
                            try: dfs_bid.append(load_investigacion_bid_en(sd, ed))
                            except: pass
                            dfs_bid = [d for d in dfs_bid if not d.empty]
                            if dfs_bid: df = pd.concat(dfs_bid, ignore_index=True).drop_duplicates(subset=['Link'])
                        elif o == "BPI": df = load_investigacion_bpi(sd, ed)
                        elif o == "BM": df = load_investigacion_bm(sd, ed)
                        
                    elif tipo_doc == "Publicaciones Institucionales":
                        if o == "BPI": df = load_pub_inst_bpi(sd, ed)
                        elif o == "CEF": df = load_pub_inst_cef(sd, ed)
                        elif o == "BM": df = load_pub_inst_bm(sd, ed)
                        elif o == "FMI": 
                            df_flagships = load_pub_inst_fmi(sd, ed)
                            df_prs = load_press_releases_fmi(sd, ed)
                            df_crs = load_country_reports_fmi(sd, ed)
                            dfs_a_unir = [d for d in [df_flagships, df_prs, df_crs] if not d.empty]
                            if dfs_a_unir:
                                df = pd.concat(dfs_a_unir, ignore_index=True)
                                df = df.sort_values("Date", ascending=False)
                        
                except Exception as e:
                    pass
                
                if not df.empty:
                    df["Date"] = pd.to_datetime(df["Date"], errors='coerce')
                    df_f = df[(df["Date"].dt.year.isin(a_num)) & (df["Date"].dt.month.isin(m_num))].copy()
                    if not df_f.empty: 
                        df_f['Organismo'] = o
                        dfs_comb.append(df_f)
                progreso.progress((i+1)/len(target_orgs))
            
            txt.empty()
            progreso.empty()
            
            if dfs_comb:
                f_df = pd.concat(dfs_comb, ignore_index=True)              
                f_df['Categoría'] = tipo_doc
                
                # --- PREPARACIÓN PARA EL WORD (Orden Institucional) ---
                if tipo_doc == "Discursos":
                    f_df_word = f_df.sort_values(by=["Title"], ascending=[True])
                else:
                    f_df_word = f_df.sort_values(by=["Organismo", "Title"], ascending=[True, True])
                    
                f_df_word = f_df_word[['Categoría', 'Organismo', 'Title', 'Link']]
                f_df_word = f_df_word.rename(columns={"Categoría": "Tipo de Documento", "Title": "Nombre de Documento"})
                
                st.success(f"Se encontraron **{len(f_df)}** documentos.")
                word_file = generate_word(f_df_word, title=f"Explorador - {tipo_doc}")
                st.download_button("📄 Descargar en Word", data=word_file, file_name=f"Explorador_{tipo_doc}.docx")
                
                # --- PREPARACIÓN PARA LA VISTA PREVIA (Orden Cronológico + Columna Fecha) ---
                disp = f_df.copy()
                disp = disp.sort_values(by="Date", ascending=False) # Orden cronológico
                disp["Fecha"] = disp["Date"].dt.strftime('%d/%m/%Y') # Formatear fecha
                disp["Nombre de Documento"] = disp.apply(lambda x: f"[{x['Title']}]({x['Link']})", axis=1)
                disp = disp.rename(columns={"Categoría": "Tipo de Documento"})
                
                # Mostrar en pantalla según si se seleccionó 'Todos' o un organismo específico
                if organismo_seleccionado == "Todos":
                    cols_vis = ["Fecha", "Tipo de Documento", "Organismo", "Nombre de Documento"] 
                else:
                    cols_vis = ["Fecha", "Tipo de Documento", "Nombre de Documento"]
                    
                st.markdown(disp[cols_vis].to_markdown(index=False), unsafe_allow_html=True)
            else: 
                st.warning("No se encontraron documentos para las fechas seleccionadas.")
